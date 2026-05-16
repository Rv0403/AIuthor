"""PDF and DOCX export with front/back matter and TOC."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from config import get_settings
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak


FRONT_MATTER_SECTIONS = [
    ("half_title", "Half-Title"),
    ("title", "Title Page"),
    ("copyright_block", "Copyright"),
    ("dedication", "Dedication"),
    ("epigraph", "Epigraph"),
    ("foreword", "Foreword"),
    ("preface", "Preface"),
    ("acknowledgments", "Acknowledgements"),
    ("introduction", "Introduction"),
]

BACK_MATTER_SECTIONS = [
    ("afterword", "Afterword"),
    ("appendix", "Appendix"),
    ("references", "References"),
    ("about_author", "About the Author"),
    ("back_cover_copy", "Back Cover Copy"),
]


def export_book(run_id: str, manifest: dict[str, Any], outline: dict[str, Any]) -> dict[str, str]:
    settings = get_settings()
    out_dir = settings.outputs_dir / run_id
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_path = out_dir / f"{_safe_name(manifest.get('title', 'book'))}.docx"
    pdf_path = out_dir / f"{_safe_name(manifest.get('title', 'book'))}.pdf"

    _build_docx(docx_path, manifest, outline)
    _build_pdf(pdf_path, manifest, outline)

    return {"docx": str(docx_path), "pdf": str(pdf_path)}


def _safe_name(title: str) -> str:
    return "".join(c if c.isalnum() or c in " -_" else "_" for c in title)[:80].strip() or "book"


def _build_docx(path: Path, manifest: dict[str, Any], outline: dict[str, Any]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    # Half-title
    ht = manifest.get("half_title") or manifest.get("title", "Untitled")
    p = doc.add_paragraph(ht)
    p.alignment = 1
    doc.add_page_break()

    # Title page
    doc.add_paragraph(manifest.get("title", "Untitled"), style="Title")
    if manifest.get("subtitle"):
        doc.add_paragraph(manifest["subtitle"])
    doc.add_page_break()

    # Copyright
    copyright_text = manifest.get("copyright_block") or _default_copyright(manifest.get("title", ""))
    doc.add_paragraph(copyright_text)
    doc.add_page_break()

    for key, label in FRONT_MATTER_SECTIONS[3:]:
        text = manifest.get(key, "")
        if text:
            doc.add_heading(label, level=1)
            doc.add_paragraph(text)
            doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for ch in outline.get("chapters", manifest.get("chapters", [])):
        num = ch.get("chapter_number", "")
        title = ch.get("title", "")
        doc.add_paragraph(f"Chapter {num}: {title}", style="List Number")
    doc.add_page_break()

    # Body — page numbers from introduction onward (simplified: section break)
    doc.add_section()
    for ch in sorted(manifest.get("chapters", []), key=lambda x: x.get("chapter_number", 0)):
        doc.add_heading(f"Chapter {ch.get('chapter_number')}: {ch.get('title', '')}", level=1)
        body = ch.get("verified_text") or ch.get("edited_text") or ch.get("humanized_text") or ch.get("raw_text", "")
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_page_break()

    # Glossary
    glossary = manifest.get("glossary", {})
    if glossary:
        doc.add_heading("Glossary", level=1)
        for term, definition in sorted(glossary.items()):
            doc.add_paragraph(f"{term}: {definition}")

    for key, label in BACK_MATTER_SECTIONS:
        if key == "references":
            refs = manifest.get("references", [])
            if refs:
                doc.add_heading(label, level=1)
                for ref in refs:
                    doc.add_paragraph(str(ref), style="List Bullet")
            continue
        text = manifest.get(key, "")
        if text and key != "back_cover_copy":
            doc.add_heading(label, level=1)
            doc.add_paragraph(text)

    if manifest.get("back_cover_copy"):
        doc.add_page_break()
        doc.add_heading("Back Cover", level=1)
        doc.add_paragraph(manifest["back_cover_copy"])

    doc.save(str(path))


def _build_pdf(path: Path, manifest: dict[str, Any], outline: dict[str, Any]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("BookTitle", parent=styles["Heading1"], fontSize=24, spaceAfter=20)
    h_style = styles["Heading2"]
    body_style = styles["BodyText"]
    story = []

    story.append(Paragraph(manifest.get("half_title") or manifest.get("title", "Untitled"), title_style))
    story.append(PageBreak())
    story.append(Paragraph(manifest.get("title", "Untitled"), title_style))
    story.append(PageBreak())
    story.append(Paragraph(_default_copyright(manifest.get("title", "")), body_style))
    story.append(PageBreak())

    for key, label in FRONT_MATTER_SECTIONS[3:]:
        text = manifest.get(key, "")
        if text:
            story.append(Paragraph(label, h_style))
            story.append(Paragraph(text.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 12))

    story.append(Paragraph("Table of Contents", h_style))
    for ch in outline.get("chapters", manifest.get("chapters", [])):
        story.append(Paragraph(f"Chapter {ch.get('chapter_number')}: {ch.get('title', '')}", body_style))
    story.append(PageBreak())

    for ch in sorted(manifest.get("chapters", []), key=lambda x: x.get("chapter_number", 0)):
        story.append(Paragraph(f"Chapter {ch.get('chapter_number')}: {ch.get('title', '')}", h_style))
        body = ch.get("verified_text") or ch.get("edited_text") or ch.get("humanized_text") or ""
        for para in body.split("\n\n")[:50]:
            if para.strip():
                safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))
                story.append(Spacer(1, 6))
        story.append(PageBreak())

    glossary = manifest.get("glossary", {})
    if glossary:
        story.append(Paragraph("Glossary", h_style))
        for term, definition in list(glossary.items())[:100]:
            story.append(Paragraph(f"<b>{term}</b>: {definition}", body_style))

    doc.build(story)


def _default_copyright(title: str) -> str:
    return f"""Copyright © 2026. All rights reserved.

ISBN: 978-0-000000-00-0 (placeholder)
Edition: First Edition

{title}
Published by AIuthor Press

CIP Data: Placeholder cataloging-in-publication record.

No part of this book may be reproduced without written permission."""
